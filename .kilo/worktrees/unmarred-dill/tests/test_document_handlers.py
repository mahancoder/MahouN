"""
Tests for Document Handlers
===========================

تست‌های جامع برای document_handlers و ocr_handler
"""

import pytest
import tempfile
import os
from pathlib import Path

# Ensure helper availability checks are defined before use in decorators
# to avoid NameError during module import time.
def _check_docx_available() -> bool:
    try:
        import docx  # type: ignore
        return True
    except Exception:
        return False


def _check_ocr_available() -> bool:
    try:
        from mahoun.pipelines.ingestion.ocr_handler import check_ocr_availability
        result = check_ocr_availability()
        return bool(result.get('is_available'))
    except Exception:
        return False


class TestDocumentHandlerAvailability:
    """Test handler availability checks"""
    
    def test_check_handler_availability(self):
        """Test that availability check returns proper structure"""
        from mahoun.pipelines.ingestion.document_handlers import check_handler_availability
        
        result = check_handler_availability()
        
        assert isinstance(result, dict)
        assert 'TxtHandler' in result
        assert 'DocxHandler' in result
        assert 'PdfHandler' in result
        assert 'ImageHandler' in result
        
        # TXT should always be available
        assert result['TxtHandler']['available'] is True
    
    def test_txt_handler_always_available(self):
        """TXT handler should always be available"""
        from mahoun.pipelines.ingestion.document_handlers import TxtHandler
        
        handler = TxtHandler()
        assert handler.available is True


class TestTxtHandler:
    """Test TXT file handling"""
    
    def test_extract_utf8_text(self):
        """Test extracting UTF-8 text"""
        from mahoun.pipelines.ingestion.document_handlers import TxtHandler
        
        handler = TxtHandler()
        
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
            f.write("این یک متن فارسی است.\nThis is English text.")
            temp_path = f.name
        
        try:
            result = handler.extract_text(temp_path)
            
            assert result.success is True
            assert "فارسی" in result.text
            assert "English" in result.text
            assert result.metadata['format'] == 'txt'
            assert result.handler_used == 'TxtHandler'
        finally:
            os.unlink(temp_path)
    
    def test_extract_persian_text(self):
        """Test extracting Persian legal text"""
        from mahoun.pipelines.ingestion.document_handlers import TxtHandler
        
        handler = TxtHandler()
        
        persian_text = """
        رأی دادگاه
        
        در خصوص دعوای آقای احمد احمدی فرزند محمد به طرفیت خانم زهرا محمدی
        با توجه به مستندات ارائه شده و ماده ۳۴۸ قانون آیین دادرسی مدنی
        حکم به رد دعوای خواهان صادر می‌گردد.
        
        این رأی قطعی است.
        """
        
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
            f.write(persian_text)
            temp_path = f.name
        
        try:
            result = handler.extract_text(temp_path)
            
            assert result.success is True
            assert "رأی دادگاه" in result.text
            assert "ماده ۳۴۸" in result.text
            assert "قطعی" in result.text
        finally:
            os.unlink(temp_path)
    
    def test_supports_txt_and_md(self):
        """Test that handler supports .txt and .md files"""
        from mahoun.pipelines.ingestion.document_handlers import TxtHandler
        
        handler = TxtHandler()
        
        assert handler.supports_file("document.txt") is True
        assert handler.supports_file("README.md") is True
        assert handler.supports_file("document.pdf") is False
        assert handler.supports_file("document.docx") is False


class TestDocxHandler:
    """Test DOCX file handling"""
    
    def test_docx_handler_init(self):
        """Test DOCX handler initialization"""
        from mahoun.pipelines.ingestion.document_handlers import DocxHandler
        
        handler = DocxHandler()
        # May or may not be available depending on python-docx installation
        assert isinstance(handler.available, bool)
    
    def test_supports_docx_only(self):
        """Test that handler only supports .docx files"""
        from mahoun.pipelines.ingestion.document_handlers import DocxHandler
        
        handler = DocxHandler()
        
        assert handler.supports_file("document.docx") is True
        assert handler.supports_file("document.doc") is False  # Old format not supported
        assert handler.supports_file("document.txt") is False
        assert handler.supports_file("document.pdf") is False
    
    @pytest.mark.skipif(
        not _check_docx_available(),
        reason="python-docx not installed"
    )
    def test_extract_docx_text(self):
        """Test extracting text from DOCX (requires python-docx)"""
        from mahoun.pipelines.ingestion.document_handlers import DocxHandler
        import docx
        
        handler = DocxHandler()
        
        # Create a test DOCX file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            temp_path = f.name
        
        try:
            doc = docx.Document()
            doc.add_heading('عنوان سند', 0)
            doc.add_paragraph('این یک پاراگراف فارسی است.')
            doc.add_paragraph('This is an English paragraph.')
            doc.save(temp_path)
            
            result = handler.extract_text(temp_path)
            
            assert result.success is True
            assert "عنوان سند" in result.text
            assert "فارسی" in result.text
            assert result.metadata['format'] == 'docx'
        finally:
            os.unlink(temp_path)


class TestPdfHandler:
    """Test PDF file handling"""
    
    def test_pdf_handler_init(self):
        """Test PDF handler initialization"""
        from mahoun.pipelines.ingestion.document_handlers import PdfHandler
        
        handler = PdfHandler()
        # Check which backends are available
        assert isinstance(handler.pdfplumber_available, bool)
        assert isinstance(handler.pypdf2_available, bool)
    
    def test_supports_pdf_only(self):
        """Test that handler only supports .pdf files"""
        from mahoun.pipelines.ingestion.document_handlers import PdfHandler
        
        handler = PdfHandler()
        
        assert handler.supports_file("document.pdf") is True
        assert handler.supports_file("document.PDF") is True
        assert handler.supports_file("document.txt") is False
        assert handler.supports_file("document.docx") is False


class TestImageHandler:
    """Test Image/OCR handling"""
    
    def test_image_handler_init(self):
        """Test Image handler initialization"""
        from mahoun.pipelines.ingestion.document_handlers import ImageHandler
        
        handler = ImageHandler()
        assert isinstance(handler.available, bool)
    
    def test_supports_image_formats(self):
        """Test that handler supports various image formats"""
        from mahoun.pipelines.ingestion.document_handlers import ImageHandler
        
        handler = ImageHandler()
        
        assert handler.supports_file("image.jpg") is True
        assert handler.supports_file("image.jpeg") is True
        assert handler.supports_file("image.png") is True
        assert handler.supports_file("image.bmp") is True
        assert handler.supports_file("image.tiff") is True
        assert handler.supports_file("image.gif") is True
        assert handler.supports_file("image.webp") is True
        assert handler.supports_file("document.pdf") is False
        assert handler.supports_file("document.txt") is False


class TestDocumentHandlerFactory:
    """Test DocumentHandlerFactory"""
    
    def test_factory_initialization(self):
        """Test factory initializes all handlers"""
        from mahoun.pipelines.ingestion.document_handlers import DocumentHandlerFactory
        
        factory = DocumentHandlerFactory()
        
        assert len(factory.handlers) >= 3  # At least TXT, DOCX, PDF
    
    def test_get_handler_for_txt(self):
        """Test getting handler for TXT file"""
        from mahoun.pipelines.ingestion.document_handlers import DocumentHandlerFactory, TxtHandler
        
        factory = DocumentHandlerFactory()
        handler = factory.get_handler("document.txt")
        
        assert handler is not None
        assert isinstance(handler, TxtHandler)
    
    def test_extract_text_convenience(self):
        """Test extract_document_text convenience function"""
        from mahoun.pipelines.ingestion.document_handlers import extract_document_text
        
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
            f.write("Test content تست محتوا")
            temp_path = f.name
        
        try:
            result = extract_document_text(temp_path)
            
            assert result.success is True
            assert "Test content" in result.text
            assert "تست" in result.text
        finally:
            os.unlink(temp_path)
    
    def test_file_not_found(self):
        """Test handling of non-existent file"""
        from mahoun.pipelines.ingestion.document_handlers import extract_document_text
        
        result = extract_document_text("/nonexistent/path/file.txt")
        
        assert result.success is False
        assert "not found" in result.error.lower()


class TestOCRHandler:
    """Test OCR functionality"""
    
    def test_ocr_availability_check(self):
        """Test OCR availability check"""
        from mahoun.pipelines.ingestion.ocr_handler import check_ocr_availability
        
        result = check_ocr_availability()
        
        assert isinstance(result, dict)
        assert 'is_available' in result
        assert 'available_engines' in result
        assert 'active_engine' in result
        assert 'recommendations' in result
    
    def test_ocr_engine_init(self):
        """Test OCR engine initialization"""
        from mahoun.pipelines.ingestion.ocr_handler import OCREngine
        
        engine = OCREngine()
        
        assert isinstance(engine.available_engines, list)
        assert isinstance(engine.is_available, bool)
    
    @pytest.mark.skipif(
        not _check_ocr_available(),
        reason="No OCR engine available"
    )
    def test_ocr_simple_image(self):
        """Test OCR on a simple image (requires OCR engine)"""
        from mahoun.pipelines.ingestion.ocr_handler import ocr_image
        from PIL import Image, ImageDraw, ImageFont
        
        # Create a simple test image with text
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as f:
            temp_path = f.name
        
        try:
            # Create image with text
            img = Image.new('RGB', (400, 100), color='white')
            draw = ImageDraw.Draw(img)
            draw.text((10, 40), "Test OCR 123", fill='black')
            img.save(temp_path)
            
            result = ocr_image(temp_path)
            
            # OCR should at least return something
            assert result.engine in ['paddle', 'tesseract', 'easyocr', 'none']
            if result.success:
                assert len(result.text) > 0
        finally:
            os.unlink(temp_path)


class TestIntegration:
    """Integration tests"""
    
    def test_full_pipeline_txt(self):
        """Test full pipeline with TXT file"""
        from mahoun.pipelines.ingestion.document_handlers import extract_document_text
        
        verdict_text = """
        بسمه تعالی
        
        رأی دادگاه تجدیدنظر استان تهران
        شعبه ۱۰
        
        در خصوص تجدیدنظرخواهی آقای احمد احمدی فرزند محمد
        از دادنامه شماره ۱۲۳۴ صادره از شعبه ۵ دادگاه عمومی حقوقی تهران
        
        با توجه به محتویات پرونده و لایحه تجدیدنظرخواهی و مستندات ابرازی
        و با عنایت به ماده ۳۴۸ و ۳۵۸ قانون آیین دادرسی مدنی
        
        نظر به اینکه تجدیدنظرخواهی با هیچ یک از جهات مندرج در ماده ۳۴۸
        قانون آیین دادرسی مدنی انطباق ندارد
        
        لذا ضمن رد تجدیدنظرخواهی، دادنامه تجدیدنظرخواسته تأیید می‌گردد.
        
        این رأی قطعی است.
        
        رئیس شعبه ۱۰ دادگاه تجدیدنظر استان تهران
        """
        
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
            f.write(verdict_text)
            temp_path = f.name
        
        try:
            result = extract_document_text(temp_path)
            
            assert result.success is True
            assert "دادگاه تجدیدنظر" in result.text
            assert "ماده ۳۴۸" in result.text
            assert "قطعی" in result.text
            assert result.metadata['format'] == 'txt'
        finally:
            os.unlink(temp_path)


# ============================================================================
# Helper Functions
# ============================================================================

def _check_docx_available() -> bool:
    """Check if python-docx is available"""
    try:
        import docx
        return True
    except ImportError:
        return False


def _check_ocr_available() -> bool:
    """Check if any OCR engine is available"""
    try:
        from mahoun.pipelines.ingestion.ocr_handler import check_ocr_availability
        result = check_ocr_availability()
        return result['is_available']
    except:
        return False


# ============================================================================
# Run Tests
# ============================================================================

if __name__ == "__main__":
    pytest.main([__file__, "-v"])
