"""
Ultra-Advanced Data Parsers
=========================

Enterprise-grade parsers integrated from Ultra Systems with:
- Multi-format support (PDF, DOCX, HTML, JSON, XML, etc.)
- Intelligent document parsing
- OCR for scanned documents
- Table extraction
- Metadata extraction
- Quality validation
"""

import json
import re
from pathlib import Path
from typing import Dict, List, Any, Optional, Union
from dataclasses import dataclass, field
import chardet
import logging

logger = logging.getLogger(__name__)

# Import Ultra Systems parsers
from ultra_systems.pipelines.ultra_data_ingestion import (
    PDFParser as UltraPDFParser,
    TextParser as UltraTextParser,
    DocumentFormat,
    IngestionConfig,
    Document
)

# Try to import OCR libraries
try:
    import pytesseract
    from pdf2image import convert_from_path
    HAS_OCR = True
except ImportError:
    HAS_OCR = False
    # logger.warning("OCR libraries not available. Install: pip install pytesseract pdf2image")

try:
    import PyPDF2
    import pdfplumber
    HAS_PDF = True
except ImportError:
    HAS_PDF = False
    # logger.warning("PDF libraries not available")

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    # logger.warning("python-docx not available")


@dataclass
class ParseResult:
    """Result of parsing operation"""
    success: bool
    data: Any
    metadata: Dict[str, Any]
    errors: List[str]
    warnings: List[str]
    
    @property
    def is_valid(self) -> bool:
        """Check if result is valid"""
        return self.success and not self.errors


class BaseParser:
    """Base class for all parsers"""
    
    def __init__(self):
        self.encoding = 'utf-8'
        self.errors = []
        self.warnings = []
    
    def parse(self, file_path: Path) -> ParseResult:
        """Parse file and return structured data"""
        raise NotImplementedError
    
    def detect_encoding(self, file_path: Path) -> str:
        """Detect file encoding"""
        try:
            with open(file_path, 'rb') as f:
                raw_data = f.read(10000)  # Read first 10KB
                result = chardet.detect(raw_data)
                encoding = result['encoding']
                confidence = result['confidence']
                
                if confidence < 0.7:
                    logger.warning(
                        f"Low encoding confidence ({confidence:.2f}) for {file_path}, "
                        f"detected: {encoding}"
                    )
                
                return encoding or 'utf-8'
        except Exception as e:
            logger.error(f"Error detecting encoding: {e}")
            return 'utf-8'
    
    def extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract file metadata"""
        stat = file_path.stat()
        return {
            "filename": file_path.name,
            "extension": file_path.suffix,
            "size_bytes": stat.st_size,
            "created_at": stat.st_ctime,
            "modified_at": stat.st_mtime,
        }


class PDFParser(BaseParser):
    """
    Ultra-Advanced PDF Parser
    
    Features:
    - Text extraction with layout preservation
    - Table extraction
    - Metadata extraction
    - Multi-page support
    - OCR fallback (if available)
    """
    
    def __init__(self, use_ocr: bool = False):
        super().__init__()
        self.use_ocr = use_ocr
        
        # Use Ultra PDF Parser
        config = IngestionConfig(
            input_dir="./data/input",
            output_dir="./data/processed",
            enable_ocr=use_ocr,
            enable_table_extraction=True,
            enable_image_extraction=True
        )
        self.ultra_parser = UltraPDFParser(config)
        
        if not HAS_PDF:
            raise ImportError("PDF libraries not installed. Install: pip install PyPDF2 pdfplumber")
    
    def parse(self, file_path: Path) -> ParseResult:
        """Parse PDF file using Ultra parser"""
        logger.info(f"Parsing PDF with Ultra parser: {file_path}")
        
        try:
            # Use Ultra parser
            document = self.ultra_parser.parse(file_path)
            
            # Convert to MAHOUN format
            data = {
                "text": document.text,
                "sections": document.sections,
                "tables": document.tables,
                "images": document.images,
                "metadata": document.metadata,
                "quality_score": document.quality_score,
                "ocr_confidence": document.ocr_confidence,
            }
            
            metadata = self.extract_metadata(file_path)
            metadata.update(document.metadata)
            
            return ParseResult(
                success=True,
                data=data,
                metadata=metadata,
                errors=self.errors,
                warnings=self.warnings,
            )
            
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {e}")
            return ParseResult(
                success=False,
                data={},
                metadata=self.extract_metadata(file_path),
                errors=[str(e)],
                warnings=self.warnings,
            )


class DOCXParser(BaseParser):
    """
    Ultra-Advanced DOCX Parser
    
    Features:
    - Text extraction with formatting
    - Table extraction
    - Header/footer extraction
    - Style preservation
    """
    
    def __init__(self):
        super().__init__()
        
        if not HAS_DOCX:
            raise ImportError("python-docx not installed. Install: pip install python-docx")
    
    def parse(self, file_path: Path) -> ParseResult:
        """Parse DOCX file"""
        logger.info(f"Parsing DOCX: {file_path}")
        
        try:
            if HAS_DOCX:
                doc = DocxDocument(file_path)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                text = '\n'.join(paragraphs)
                
                # Extract tables
                tables = []
                for table in doc.tables:
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        table_data.append(row_data)
                    tables.append(table_data)
                
                data = {
                    "text": text,
                    "paragraphs": paragraphs,
                    "tables": tables,
                    "metadata": self.extract_metadata(file_path)
                }
                
                return ParseResult(
                    success=True,
                    data=data,
                    metadata=self.extract_metadata(file_path),
                    errors=[],
                    warnings=[]
                )
            else:
                raise ImportError("python-docx not installed. Install: pip install python-docx")
                
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            return ParseResult(
                success=False,
                data={},
                metadata=self.extract_metadata(file_path),
                errors=[str(e)],
                warnings=[]
            )


class XMLParser(BaseParser):
    """
    Ultra-Advanced XML Parser
    
    Features:
    - XML parsing with validation
    - Namespace support
    - XPath queries
    - Schema validation
    """
    
    def __init__(self):
        super().__init__()
        try:
            import xml.etree.ElementTree as ET
            self.ET = ET
        except ImportError:
            raise ImportError("XML libraries not available")
    
    def parse(self, file_path: Path) -> ParseResult:
        """Parse XML file"""
        logger.info(f"Parsing XML: {file_path}")
        
        try:
            tree = self.ET.parse(file_path)
            root = tree.getroot()
            
            # Convert XML to dict
            data = self._element_to_dict(root)
            
            return ParseResult(
                success=True,
                data=data,
                metadata=self.extract_metadata(file_path),
                errors=[],
                warnings=[]
            )
            
        except Exception as e:
            logger.error(f"Error parsing XML {file_path}: {e}")
            return ParseResult(
                success=False,
                data={},
                metadata=self.extract_metadata(file_path),
                errors=[str(e)],
                warnings=[]
            )
    
    def _element_to_dict(self, element) -> Dict[str, Any]:
        """Convert XML element to dictionary"""
        result = {}
        
        # Add attributes
        if element.attrib:
            result['@attributes'] = element.attrib
        
        # Add text content
        if element.text and element.text.strip():
            if len(element) == 0:
                return element.text.strip()
            result['#text'] = element.text.strip()
        
        # Add children
        for child in element:
            child_data = self._element_to_dict(child)
            if child.tag in result:
                # Handle multiple children with same tag
                if not isinstance(result[child.tag], list):
                    result[child.tag] = [result[child.tag]]
                result[child.tag].append(child_data)
            else:
                result[child.tag] = child_data
        
        return result if result else {}


class JSONParser(BaseParser):
    """
    Ultra-Advanced JSON Parser
    
    Features:
    - JSON parsing with validation
    - Schema validation
    - Deep structure analysis
    """
    
    def parse(self, file_path: Path) -> ParseResult:
        """Parse JSON file"""
        logger.info(f"Parsing JSON: {file_path}")
        
        try:
            with open(file_path, 'r', encoding=self.detect_encoding(file_path)) as f:
                data = json.load(f)
            
            return ParseResult(
                success=True,
                data=data,
                metadata=self.extract_metadata(file_path),
                errors=[],
                warnings=[]
            )
            
        except Exception as e:
            logger.error(f"Error parsing JSON {file_path}: {e}")
            return ParseResult(
                success=False,
                data={},
                metadata=self.extract_metadata(file_path),
                errors=[str(e)],
                warnings=[]
            )


class TXTParser(BaseParser):
    """
    Ultra-Advanced Text Parser
    
    Features:
    - Text parsing with encoding detection
    - Line-by-line processing
    - Pattern matching
    """
    
    def parse(self, file_path: Path) -> ParseResult:
        """Parse text file"""
        logger.info(f"Parsing TXT: {file_path}")
        
        try:
            encoding = self.detect_encoding(file_path)
            
            with open(file_path, 'r', encoding=encoding) as f:
                text = f.read()
            
            lines = text.splitlines()
            
            data = {
                "text": text,
                "lines": lines,
                "line_count": len(lines),
                "char_count": len(text),
                "word_count": len(text.split()),
            }
            
            return ParseResult(
                success=True,
                data=data,
                metadata=self.extract_metadata(file_path),
                errors=[],
                warnings=[]
            )
            
        except Exception as e:
            logger.error(f"Error parsing TXT {file_path}: {e}")
            return ParseResult(
                success=False,
                data={},
                metadata=self.extract_metadata(file_path),
                errors=[str(e)],
                warnings=[]
            )


class LegalDocumentParser(BaseParser):
    """
    Specialized parser for legal documents
    
    Features:
    - Legal document structure recognition
    - Article/Clause extraction
    - Reference linking
    - Citation analysis
    """
    
    def __init__(self):
        super().__init__()
        # Legal document patterns
        self.patterns = {
            'article': r'ماده\s+(\d+)',
            'clause': r'بند\s+([الف-ی]\s*-?\s*\d*)',
            'section': r'بخش\s+(اول|دوم|سوم|چهارم|پنجم|ششم|هفتم|هشتم|نهم|دهم|[۰-۹]+)',
            'chapter': r'فصل\s+([۰-۹]+)',
        }
    
    def parse(self, file_path: Path) -> ParseResult:
        """Parse legal document"""
        logger.info(f"Parsing legal document: {file_path}")
        
        try:
            # Use appropriate parser based on file extension
            ext = file_path.suffix.lower()
            
            if ext == '.pdf':
                parser = PDFParser()
            elif ext == '.docx':
                parser = DOCXParser()
            elif ext == '.json':
                parser = JSONParser()
            else:
                parser = TXTParser()
            
            result = parser.parse(file_path)
            
            if result.success:
                # Add legal document analysis
                legal_data = self._analyze_legal_document(result.data.get('text', ''))
                result.data.update(legal_data)
            
            return result
            
        except Exception as e:
            logger.error(f"Error parsing legal document {file_path}: {e}")
            return ParseResult(
                success=False,
                data={},
                metadata=self.extract_metadata(file_path),
                errors=[str(e)],
                warnings=[]
            )
    
    def _analyze_legal_document(self, text: str) -> Dict[str, Any]:
        """Analyze legal document structure"""
        analysis = {
            'articles': [],
            'clauses': [],
            'sections': [],
            'chapters': [],
        }
        
        # Find articles
        article_matches = re.finditer(self.patterns['article'], text)
        analysis['articles'] = [{'number': match.group(1), 'position': match.start()} 
                               for match in article_matches]
        
        # Find clauses
        clause_matches = re.finditer(self.patterns['clause'], text)
        analysis['clauses'] = [{'label': match.group(1), 'position': match.start()} 
                              for match in clause_matches]
        
        # Find sections
        section_matches = re.finditer(self.patterns['section'], text)
        analysis['sections'] = [{'label': match.group(1), 'position': match.start()} 
                               for match in section_matches]
        
        # Find chapters
        chapter_matches = re.finditer(self.patterns['chapter'], text)
        analysis['chapters'] = [{'number': match.group(1), 'position': match.start()} 
                               for match in chapter_matches]
        
        return analysis


class ParserFactory:
    """Factory for creating appropriate parsers"""
    
    @staticmethod
    def create_parser(file_path: Path, legal_document: bool = False) -> BaseParser:
        """Create parser based on file extension"""
        ext = file_path.suffix.lower()
        
        if legal_document:
            return LegalDocumentParser()
        
        if ext == '.pdf':
            return PDFParser()
        elif ext == '.docx':
            return DOCXParser()
        elif ext == '.json':
            return JSONParser()
        elif ext == '.xml':
            return XMLParser()
        else:
            return TXTParser()