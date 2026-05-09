"""
Document Format Handlers for MAHOUN Ingestion Pipeline
=======================================================

Loosely-coupled document format handlers with graceful degradation.
Supports: TXT, DOCX, PDF (with OCR fallback), Images

Design Principles:
- Fail gracefully if dependencies missing
- Never crash orchestrator initialization
- Return structured text output
- Maintain Desktop-Minimal compatibility

Dependencies (optional):
- python-docx: DOCX support
- pypdf: Basic PDF text extraction
- pdfplumber: Advanced PDF extraction (tables, layout)
- pdf2image + poppler: PDF to image conversion for OCR
- pytesseract + tesseract-ocr: OCR engine
- paddleocr: Alternative OCR (better for Persian)
"""

import logging
from typing import Any, Dict, List, Optional
from pathlib import Path
from dataclasses import dataclass
import tempfile
import os

logger = logging.getLogger(__name__)


@dataclass
class DocumentExtractionResult:
    """Result of document text extraction"""
    success: bool
    text: str
    metadata: Dict[str, Any]
    error: Optional[str] = None
    handler_used: str = "unknown"


class BaseDocumentHandler:
    """Base class for document handlers"""
    
    def __init__(self):
        self.available = self._check_dependencies()
        if not self.available:
            logger.warning(f"{self.__class__.__name__} dependencies not available")
    
    def _check_dependencies(self) -> bool:
        """Check if required dependencies are available"""
        raise NotImplementedError
    
    def extract_text(self, file_path: str) -> DocumentExtractionResult:
        """Extract text from document"""
        raise NotImplementedError
    
    def supports_file(self, file_path: str) -> bool:
        """Check if this handler supports the file"""
        raise NotImplementedError


class TxtHandler(BaseDocumentHandler):
    """Plain text file handler (always available)"""
    
    def _check_dependencies(self) -> bool:
        return True  # No dependencies needed
    
    def supports_file(self, file_path: str) -> bool:
        return file_path.lower().endswith(('.txt', '.md'))
    
    def extract_text(self, file_path: str) -> DocumentExtractionResult:
        """Extract text from TXT file"""
        try:
            # Try multiple encodings
            encodings = ['utf-8', 'utf-8-sig', 'cp1256', 'windows-1256', 'iso-8859-1']
            text: Optional[Any] = None
            used_encoding: Optional[Any] = None
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    used_encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                # Last resort: read as binary and decode with errors='replace'
                with open(file_path, 'rb') as f:
                    text = f.read().decode('utf-8', errors='replace')
                used_encoding = 'utf-8 (with replacements)'
            
            return DocumentExtractionResult(
                success=True,
                text=text,
                metadata={
                    "format": "txt",
                    "file_path": file_path,
                    "file_size": Path(file_path).stat().st_size,
                    "encoding": used_encoding,
                    "extraction_method": "native"
                },
                handler_used="TxtHandler"
            )
        except Exception as e:
            logger.error(f"Failed to read TXT file {file_path}: {e}")
            return DocumentExtractionResult(
                success=False,
                text="",
                metadata={"format": "txt", "file_path": file_path},
                error=str(e),
                handler_used="TxtHandler"
            )


class DocxHandler(BaseDocumentHandler):
    """DOCX file handler using python-docx"""
    
    def _check_dependencies(self) -> bool:
        try:
            import docx
            return True
        except ImportError:
            return False
    
    def supports_file(self, file_path: str) -> bool:
        return file_path.lower().endswith('.docx')
    
    def extract_text(self, file_path: str) -> DocumentExtractionResult:
        """Extract text from DOCX file with full structure"""
        if not self.available:
            return DocumentExtractionResult(
                success=False,
                text="",
                metadata={"format": "docx", "file_path": file_path},
                error="python-docx library not available. Install: pip install python-docx",
                handler_used="DocxHandler"
            )
        
        try:
            import docx
            
            doc = docx.Document(file_path)
            
            # Extract paragraphs with style info
            paragraphs: List[Any] = []
            for p in doc.paragraphs:
                if p.text.strip():
                    # Check if heading
                    if p.style and p.style.name.startswith('Heading'):
                        paragraphs.append(f"\n{'#' * int(p.style.name[-1]) if p.style.name[-1].isdigit() else '##'} {p.text}\n")
                    else:
                        paragraphs.append(p.text)
            
            # Extract tables
            table_texts: List[Any] = []
            for table_idx, table in enumerate(doc.tables):
                table_rows: List[Any] = []
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_rows.append(row_text)
                if table_rows:
                    table_texts.append(f"\n=== جدول {table_idx + 1} ===\n" + '\n'.join(table_rows))
            
            # Extract headers and footers
            headers: List[Any] = []
            footers: List[Any] = []
            for section in doc.sections:
                if section.header and section.header.paragraphs:
                    for p in section.header.paragraphs:
                        if p.text.strip():
                            headers.append(p.text.strip())
                if section.footer and section.footer.paragraphs:
                    for p in section.footer.paragraphs:
                        if p.text.strip():
                            footers.append(p.text.strip())
            
            # Combine all text
            all_parts: List[Any] = []
            if headers:
                all_parts.append("=== سربرگ ===\n" + '\n'.join(set(headers)))
            all_parts.append('\n\n'.join(paragraphs))
            if table_texts:
                all_parts.extend(table_texts)
            if footers:
                all_parts.append("\n=== پاورقی ===\n" + '\n'.join(set(footers)))
            
            all_text = '\n\n'.join(all_parts)
            
            return DocumentExtractionResult(
                success=True,
                text=all_text,
                metadata={
                    "format": "docx",
                    "file_path": file_path,
                    "file_size": Path(file_path).stat().st_size,
                    "num_paragraphs": len(paragraphs),
                    "num_tables": len(doc.tables),
                    "has_headers": len(headers) > 0,
                    "has_footers": len(footers) > 0,
                    "extraction_method": "python-docx"
                },
                handler_used="DocxHandler"
            )
            
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
            return DocumentExtractionResult(
                success=False,
                text="",
                metadata={"format": "docx", "file_path": file_path},
                error=str(e),
                handler_used="DocxHandler"
            )


class PdfHandler(BaseDocumentHandler):
    """
    Advanced PDF handler with multiple extraction strategies:
    1. pdfplumber (best for tables and layout)
    2. pypdf (fallback for simple PDFs)
    3. OCR (for scanned PDFs)
    """
    
    def __init__(self):
        self.pdfplumber_available = False
        self.pypdf_available = False
        self.ocr_available = False
        self.pdf2image_available = False
        
        # Check pdfplumber (preferred)
        try:
            import pdfplumber
            self.pdfplumber_available = True
        except ImportError:
            pass
        
        # Check pypdf (fallback)
        try:
            import pypdf
            self.pypdf_available = True
        except ImportError:
            pass
        
        # Check pdf2image for OCR
        try:
            from pdf2image import convert_from_path
            self.pdf2image_available = True
        except ImportError:
            pass
        
        # Check OCR availability
        try:
            import pytesseract
            self.ocr_available = True
        except ImportError:
            try:
                from paddleocr import PaddleOCR
                self.ocr_available = True
            except ImportError:
                pass
        
        super().__init__()
    
    def _check_dependencies(self) -> bool:
        return self.pdfplumber_available or self.pypdf_available
    
    def supports_file(self, file_path: str) -> bool:
        return file_path.lower().endswith('.pdf')
    
    def extract_text(self, file_path: str) -> DocumentExtractionResult:
        """Extract text from PDF with multiple strategies"""
        
        # Strategy 1: Try pdfplumber (best for tables)
        if self.pdfplumber_available:
            result = self._extract_with_pdfplumber(file_path)
            if result.success and len(result.text.strip()) > 100:
                return result
        
        # Strategy 2: Try pypdf
        if self.pypdf_available:
            result = self._extract_with_pypdf(file_path)
            if result.success and len(result.text.strip()) > 100:
                return result
        
        # Strategy 3: OCR fallback for scanned PDFs
        if self.pdf2image_available and self.ocr_available:
            logger.info(f"Text extraction failed, trying OCR for: {file_path}")
            result = self._extract_with_ocr(file_path)
            if result.success:
                return result
        
        # No method worked
        return DocumentExtractionResult(
            success=False,
            text="",
            metadata={"format": "pdf", "file_path": file_path},
            error="Could not extract text from PDF. Install: pip install pypdf",
            handler_used="PdfHandler"
        )
    
    def _extract_with_pdfplumber(self, file_path: str) -> DocumentExtractionResult:
        """Extract using pdfplumber (best for tables)"""
        try:
            import pdfplumber
            
            pages_text: List[Any] = []
            tables_text: List[Any] = []
            with pdfplumber.open(file_path) as pdf:
                num_pages = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages):
                    # Extract text
                    text = page.extract_text()
                    if text and text.strip():
                        pages_text.append(f"=== صفحه {page_num + 1} ===\n{text}")
                    
                    # Extract tables
                    tables = page.extract_tables()
                    for table_idx, table in enumerate(tables):
                        if table:
                            table_rows: List[Any] = []
                            for row in table:
                                row_text = ' | '.join(str(cell or '').strip() for cell in row)
                                if row_text.strip():
                                    table_rows.append(row_text)
                            if table_rows:
                                tables_text.append(
                                    f"\n=== جدول صفحه {page_num + 1} ===\n" + 
                                    '\n'.join(table_rows)
                                )
            
            all_text = '\n\n'.join(pages_text)
            if tables_text:
                all_text += '\n\n' + '\n'.join(tables_text)
            
            return DocumentExtractionResult(
                success=True,
                text=all_text,
                metadata={
                    "format": "pdf",
                    "file_path": file_path,
                    "file_size": Path(file_path).stat().st_size,
                    "num_pages": num_pages,
                    "num_tables": len(tables_text),
                    "extraction_method": "pdfplumber"
                },
                handler_used="PdfHandler"
            )
            
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {e}")
            return DocumentExtractionResult(
                success=False,
                text="",
                metadata={"format": "pdf", "file_path": file_path},
                error=str(e),
                handler_used="PdfHandler"
            )
    
    def _extract_with_pypdf(self, file_path: str) -> DocumentExtractionResult:
        """Extract using pypdf (simple PDFs)"""
        try:
            from pypdf import PdfReader
            
            with open(file_path, 'rb') as f:
                pdf_reader = PdfReader(f)
                num_pages = len(pdf_reader.pages)
                
                pages_text: List[Any] = []
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        pages_text.append(f"=== صفحه {page_num + 1} ===\n{page_text}")
                
                all_text = '\n\n'.join(pages_text)
                
                return DocumentExtractionResult(
                    success=True,
                    text=all_text,
                    metadata={
                        "format": "pdf",
                        "file_path": file_path,
                        "file_size": Path(file_path).stat().st_size,
                        "num_pages": num_pages,
                        "extraction_method": "pypdf"
                    },
                    handler_used="PdfHandler"
                )
                
        except Exception as e:
            logger.warning(f"pypdf extraction failed: {e}")
            return DocumentExtractionResult(
                success=False,
                text="",
                metadata={"format": "pdf", "file_path": file_path},
                error=str(e),
                handler_used="PdfHandler"
            )
    
    def _extract_with_ocr(self, file_path: str) -> DocumentExtractionResult:
        """Extract using OCR (for scanned PDFs)"""
        try:
            from pdf2image import convert_from_path
            
            # Convert PDF pages to images
            with tempfile.TemporaryDirectory() as temp_dir:
                images = convert_from_path(
                    file_path,
                    dpi=300,
                    output_folder=temp_dir,
                    fmt='png'
                )
                
                pages_text: List[Any] = []
                # Try PaddleOCR first (better for Persian)
                try:
                    from paddleocr import PaddleOCR
                    paddle_ocr = PaddleOCR(use_angle_cls=True, lang='fa')
                    ocr_engine_name = "paddleocr"
                    
                    for page_num, image in enumerate(images):
                        # Save image temporarily
                        img_path = os.path.join(temp_dir, f"page_{page_num}.png")
                        image.save(img_path)
                        
                        # OCR
                        result = paddle_ocr.ocr(img_path, cls=True)
                        
                        if result and result[0]:
                            text_parts: List[Any] = []
                            for line in result[0]:
                                if line and len(line) >= 2:
                                    text = line[1][0] if isinstance(line[1], (list, tuple)) else line[1]
                                    if text:
                                        text_parts.append(text)
                            
                            if text_parts:
                                pages_text.append(f"=== صفحه {page_num + 1} ===\n" + '\n'.join(text_parts))
                
                except ImportError:
                    # Fallback to Tesseract
                    import pytesseract
                    ocr_engine_name = "tesseract"
                    
                    for page_num, image in enumerate(images):
                        # OCR with Persian + English
                        text = pytesseract.image_to_string(image, lang='fas+eng')
                        if text and text.strip():
                            pages_text.append(f"=== صفحه {page_num + 1} ===\n{text}")
                
                all_text = '\n\n'.join(pages_text)
                
                return DocumentExtractionResult(
                    success=len(all_text.strip()) > 0,
                    text=all_text,
                    metadata={
                        "format": "pdf",
                        "file_path": file_path,
                        "file_size": Path(file_path).stat().st_size,
                        "num_pages": len(images),
                        "extraction_method": f"ocr_{ocr_engine_name}",
                        "is_scanned": True
                    },
                    handler_used="PdfHandler"
                )
                
        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            return DocumentExtractionResult(
                success=False,
                text="",
                metadata={"format": "pdf", "file_path": file_path},
                error=f"OCR failed: {e}",
                handler_used="PdfHandler"
            )


class ImageHandler(BaseDocumentHandler):
    """
    Image handler with OCR support
    Supports: JPG, PNG, BMP, TIFF, GIF
    """
    
    def __init__(self):
        # LAZY INITIALIZATION - No OCR engines initialized here
        self.paddle_available = None  # None means not checked yet
        self.tesseract_available = None  # None means not checked yet
        self.ocr_engine = None
        self.paddle_ocr = None
        
        super().__init__()
    
    def _lazy_init_ocr(self):
        """Lazy initialization of OCR engines - only when needed"""
        # Only check if not already determined
        if self.paddle_available is None and self.tesseract_available is None:
            # Try PaddleOCR first (better for Persian)
            try:
                from paddleocr import PaddleOCR
                self.paddle_ocr = PaddleOCR(use_angle_cls=True, lang='fa')
                self.paddle_available = True
                self.ocr_engine = "paddleocr"
            except ImportError:
                self.paddle_ocr = None
                self.paddle_available = False
            
            # Fallback to Tesseract only if PaddleOCR not available
            if not self.paddle_available:
                try:
                    import pytesseract
                    from PIL import Image
                    self.tesseract_available = True
                    self.ocr_engine = "tesseract"
                except ImportError:
                    self.tesseract_available = False
            
            # Update availability
            self.available = self.paddle_available or self.tesseract_available
    
    def _check_dependencies(self) -> bool:
        # For ImageHandler, we don't want to trigger OCR initialization during dependency check
        # We only check if we can potentially support images (have required base dependencies)
        # Actual OCR engine availability is checked in _lazy_init_ocr when needed
        return True  # ImageHandler can potentially work, actual availability checked lazily
    
    def supports_file(self, file_path: str) -> bool:
        image_extensions = ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif', '.gif', '.webp']
        return any(file_path.lower().endswith(ext) for ext in image_extensions)
    
    def extract_text(self, file_path: str) -> DocumentExtractionResult:
        """Extract text from image using OCR"""
        # Lazy initialize OCR engines when actually needed
        self._lazy_init_ocr()
        
        if not self.available:
            return DocumentExtractionResult(
                success=False,
                text="",
                metadata={"format": "image", "file_path": file_path},
                error="OCR not available. Install: pip install paddleocr or pip install pytesseract",
                handler_used="ImageHandler"
            )
        
        try:
            if self.paddle_available:
                return self._extract_with_paddle(file_path)
            else:
                return self._extract_with_tesseract(file_path)
        except Exception as e:
            logger.error(f"OCR extraction failed for {file_path}: {e}")
            return DocumentExtractionResult(
                success=False,
                text="",
                metadata={"format": "image", "file_path": file_path},
                error=str(e),
                handler_used="ImageHandler"
            )
    
    def _extract_with_paddle(self, file_path: str) -> DocumentExtractionResult:
        """Extract text using PaddleOCR"""
        # Ensure OCR is initialized (should be already but double-check)
        if self.paddle_ocr is None:
            from paddleocr import PaddleOCR
            self.paddle_ocr = PaddleOCR(use_angle_cls=True, lang='fa')
        
        result = self.paddle_ocr.ocr(file_path, cls=True)
        
        text_parts: List[Any] = []
        confidences: List[Any] = []
        if result and result[0]:
            for line in result[0]:
                if line and len(line) >= 2:
                    text = line[1][0] if isinstance(line[1], (list, tuple)) else line[1]
                    conf = line[1][1] if isinstance(line[1], (list, tuple)) and len(line[1]) > 1 else 1.0
                    if text:
                        text_parts.append(text)
                        confidences.append(conf)
        
        text = '\n'.join(text_parts)
        avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0
        
        return DocumentExtractionResult(
            success=len(text.strip()) > 0,
            text=text,
            metadata={
                "format": "image",
                "file_path": file_path,
                "file_size": Path(file_path).stat().st_size,
                "ocr_engine": "paddleocr",
                "lines_extracted": len(text_parts),
                "avg_confidence": round(avg_confidence, 3)
            },
            handler_used="ImageHandler"
        )
    
    def _extract_with_tesseract(self, file_path: str) -> DocumentExtractionResult:
        """Extract text using Tesseract"""
        import pytesseract
        from PIL import Image
        
        image = Image.open(file_path)
        
        # Get image info
        width, height = image.size
        
        # Extract text (Persian + English)
        text = pytesseract.image_to_string(image, lang='fas+eng')
        
        # Get confidence data
        data = pytesseract.image_to_data(image, lang='fas+eng', output_type=pytesseract.Output.DICT)
        confidences = [int(c) for c in data['conf'] if int(c) > 0]
        avg_confidence = sum(confidences) / len(confidences) / 100 if confidences else 0.0
        
        return DocumentExtractionResult(
            success=len(text.strip()) > 0,
            text=text,
            metadata={
                "format": "image",
                "file_path": file_path,
                "file_size": Path(file_path).stat().st_size,
                "ocr_engine": "tesseract",
                "image_size": (width, height),
                "avg_confidence": round(avg_confidence, 3)
            },
            handler_used="ImageHandler"
        )


class DocumentHandlerFactory:
    """
    Factory for selecting appropriate document handler.
    
    Implements graceful fallback strategy:
    1. Try format-specific handler (DOCX/PDF/Image)
    2. Fall back to TXT handler if dependencies missing
    3. Never crash - always return a result
    """
    
    def __init__(self):
        """Initialize handlers lazily - only TxtHandler and DocxHandler instantiated here"""
        # Always instantiate TxtHandler (no dependencies)
        self.txt_handler = TxtHandler()
        
        # Instantiate DocxHandler (dependency check happens in constructor)
        self.docx_handler = DocxHandler()
        
        # Instantiate PdfHandler (dependency check happens in constructor)
        self.pdf_handler = PdfHandler()
        
        # Don't instantiate ImageHandler eagerly - only when needed
        self.image_handler = None
        
        # Keep list of handlers for iteration
        self.handlers = [self.txt_handler, self.docx_handler, self.pdf_handler]
        
        # Log availability
        available = [h.__class__.__name__ for h in self.handlers if h.available]
        unavailable = [h.__class__.__name__ for h in self.handlers if not h.available]
        
        logger.info(f"Document handlers available: {available}")
        if unavailable:
            logger.warning(f"Document handlers unavailable: {unavailable}")
    
    def _get_image_handler(self):
        """Lazy initialization of ImageHandler"""
        if self.image_handler is None:
            self.image_handler = ImageHandler()
        return self.image_handler
    
    def get_handler(self, file_path: str) -> Optional[BaseDocumentHandler]:
        """Get appropriate handler for file"""
        # Check non-image handlers first
        for handler in self.handlers:
            if handler.supports_file(file_path) and handler.available:
                return handler
        
        # Only instantiate ImageHandler if file is an image
        image_extensions = ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif', '.gif', '.webp']
        if any(file_path.lower().endswith(ext) for ext in image_extensions):
            image_handler = self._get_image_handler()
            if image_handler.supports_file(file_path) and image_handler.available:
                return image_handler
                
        return None
    
    def extract_text(self, file_path: str) -> DocumentExtractionResult:
        """
        Extract text from document with graceful fallback.
        
        Args:
            file_path: Path to document file
            
        Returns:
            DocumentExtractionResult with extracted text
        """
        if not Path(file_path).exists():
            return DocumentExtractionResult(
                success=False,
                text="",
                metadata={"file_path": file_path},
                error=f"File not found: {file_path}",
                handler_used="DocumentHandlerFactory"
            )
        
        # Try to get specific handler
        handler = self.get_handler(file_path)
        
        if handler:
            logger.debug(f"Using {handler.__class__.__name__} for {file_path}")
            return handler.extract_text(file_path)
        
        # Graceful fallback: try TXT handler
        logger.warning(
            f"No specific handler available for {file_path}, "
            f"attempting TXT fallback"
        )
        txt_handler = TxtHandler()
        result = txt_handler.extract_text(file_path)
        result.metadata["fallback_used"] = True
        return result


# ============================================================================
# Convenience Functions
# ============================================================================

def extract_document_text(file_path: str) -> DocumentExtractionResult:
    """
    Convenience function to extract text from any supported document format.
    
    Supported formats:
    - TXT, MD: Plain text
    - DOCX: Microsoft Word (requires python-docx)
    - PDF: Portable Document Format (requires pdfplumber or pypdf)
    - Images (JPG, PNG, etc.): OCR (requires paddleocr or pytesseract)
    
    Args:
        file_path: Path to document file
        
    Returns:
        DocumentExtractionResult with extracted text and metadata
        
    Example:
        >>> result = extract_document_text("verdict.pdf")
        >>> if result.success:
        ...     print(result.text)
        ...     print(f"Used: {result.handler_used}")
        ...     print(f"Method: {result.metadata.get('extraction_method')}")
    """
    factory = DocumentHandlerFactory()
    return factory.extract_text(file_path)


def check_handler_availability() -> Dict[str, Dict[str, Any]]:
    """
    Check which document handlers are available and their capabilities.
    
    Returns:
        Dictionary mapping handler name to availability and details
        
    Example:
        >>> availability = check_handler_availability()
        >>> print(f"PDF support: {availability['PdfHandler']['available']}")
        >>> print(f"PDF methods: {availability['PdfHandler']['methods']}")
    """
    result: Dict[str, Any] = {}
    # TXT Handler
    result['TxtHandler'] = {
        'available': True,
        'formats': ['.txt', '.md'],
        'methods': ['native']
    }
    
    # DOCX Handler
    try:
        import docx
        result['DocxHandler'] = {
            'available': True,
            'formats': ['.docx'],
            'methods': ['python-docx']
        }
    except ImportError:
        result['DocxHandler'] = {
            'available': False,
            'formats': ['.docx'],
            'install': 'pip install python-docx'
        }
    
    # PDF Handler
    pdf_methods: List[Any] = []
    try:
        import pdfplumber
        pdf_methods.append('pdfplumber')
    except ImportError:
        pass
    try:
        import pypdf
        pdf_methods.append('pypdf')
    except ImportError:
        pass
    try:
        from pdf2image import convert_from_path
        pdf_methods.append('pdf2image')
    except ImportError:
        pass
    
    result['PdfHandler'] = {
        'available': len(pdf_methods) > 0,
        'formats': ['.pdf'],
        'methods': pdf_methods if pdf_methods else None,
        'install': 'pip install pdfplumber pypdf pdf2image' if not pdf_methods else None
    }
    
    # Image/OCR Handler
    ocr_methods: List[Any] = []
    try:
        from paddleocr import PaddleOCR
        ocr_methods.append('paddleocr')
    except ImportError:
        pass
    try:
        import pytesseract
        ocr_methods.append('tesseract')
    except ImportError:
        pass
    
    result['ImageHandler'] = {
        'available': len(ocr_methods) > 0,
        'formats': ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.gif'],
        'methods': ocr_methods if ocr_methods else None,
        'install': 'pip install paddleocr or pip install pytesseract' if not ocr_methods else None
    }
    
    return result
